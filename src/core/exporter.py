"""
ATS-GOD Exporter
Generates downloadable DOCX and TXT reports from optimization results.
"""
import io
from datetime import datetime
from typing import Dict, Any, Optional


def export_to_docx(results: Dict[str, Any], variant: str = "balanced") -> Optional[bytes]:
    """Export full results as a formatted Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('ATS-GOD Optimization Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        summary = results.get("summary", {})
        score = summary.get("overall_score", 0)

        # Meta
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
        doc.add_paragraph(f"AI Provider: {results.get('llm_provider', 'Rule-Based')} — {results.get('llm_model', 'N/A')}")
        doc.add_paragraph(f"Overall Score: {score}/100")
        doc.add_paragraph(f"Interview Probability: {summary.get('interview_probability', 0)}%")
        doc.add_paragraph(f"Recommended Variant: {summary.get('recommended_variant', 'BALANCED')}")
        doc.add_paragraph(f"Verdict: {summary.get('verdict', '')}")

        # Agent scores
        doc.add_heading('Agent Scores', level=1)
        for agent, s in summary.get('agent_scores', {}).items():
            p = doc.add_paragraph()
            r = p.add_run(f"{agent.replace('_', ' ').title()}: {s}/100")
            if s >= 80:
                r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            elif s >= 60:
                r.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
            else:
                r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

        # Action items
        doc.add_heading('Priority Action Items', level=1)
        for i, item in enumerate(results.get('action_items', [])[:15], 1):
            doc.add_paragraph(f"{i}. {item}", style='List Number')

        # CV Variant
        variants = results.get('cv_variants', {})
        cv_content = variants.get(variant, variants.get('balanced', ''))
        if cv_content:
            doc.add_page_break()
            doc.add_heading(f'CV Variant: {variant.upper().replace("_", "-")}', level=1)
            for line in cv_content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    if line.startswith('═') or line.startswith('╔') or line.startswith('║'):
                        p.runs[0].bold = True

        # Cover letter
        cl = results.get('cover_letter', '')
        if cl:
            doc.add_page_break()
            doc.add_heading('Cover Letter', level=1)
            for para in cl.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Detailed agent reports
        doc.add_page_break()
        doc.add_heading('Detailed Agent Reports', level=1)
        for name, data in results.get('agent_results', {}).items():
            doc.add_heading(f"{name.replace('_', ' ').title()} — {data.get('score', 0)}/100", level=2)
            doc.add_paragraph("Findings:")
            for f in data.get('findings', []):
                doc.add_paragraph(f"• {f}")
            doc.add_paragraph("Recommendations:")
            for r in data.get('recommendations', []):
                doc.add_paragraph(f"→ {r}")
            if data.get('optimized_content'):
                doc.add_paragraph("AI-Generated Improvement:")
                doc.add_paragraph(data['optimized_content'])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except ImportError:
        return None
    except Exception as e:
        return None


def export_to_txt(results: Dict[str, Any]) -> str:
    """Export as plain text — always works, no dependencies needed."""

    summary = results.get("summary", {})
    score = summary.get("overall_score", 0)
    sep = "=" * 65

    lines = [
        sep,
        "ATS-GOD OPTIMIZATION REPORT",
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
        f"AI Provider: {results.get('llm_provider', 'Rule-Based')} — {results.get('llm_model', 'N/A')}",
        sep,
        "",
        f"OVERALL SCORE:         {score}/100",
        f"INTERVIEW PROBABILITY: {summary.get('interview_probability', 0)}%",
        f"RECOMMENDED VARIANT:   {summary.get('recommended_variant', 'BALANCED')}",
        f"VERDICT:               {summary.get('verdict', '')}",
        f"WEAKEST AREA:          {summary.get('weakest_area', '').replace('_', ' ').title()}",
        f"STRONGEST AREA:        {summary.get('strongest_area', '').replace('_', ' ').title()}",
        "",
        sep,
        "AGENT SCORES",
        sep,
    ]

    icons = {
        "algorithm_breaker": "🎯", "sa_specialist": "🇿🇦", "global_setter": "🌍",
        "recruiter_scanner": "👁️", "hiring_manager": "💼", "semantic_matcher": "📊",
        "compliance_guardian": "⚖️", "future_architect": "🚀",
    }

    for name, s in summary.get('agent_scores', {}).items():
        rating = "STRONG ✓" if s >= 80 else "NEEDS WORK !" if s < 60 else "ADEQUATE ~"
        icon = icons.get(name, "•")
        lines.append(f"  {icon} {name.replace('_', ' ').title():<30} {s:>3}/100  {rating}")

    lines += ["", sep, "PRIORITY ACTION ITEMS", sep, ""]
    for i, item in enumerate(results.get('action_items', [])[:15], 1):
        priority = "🔴 URGENT" if i <= 4 else "🟡 IMPORTANT" if i <= 8 else "🟢 HELPFUL"
        lines.append(f"  {i:>2}. [{priority}] {item}")

    for variant_key, variant_label in [
        ("balanced", "BALANCED VARIANT ⭐ RECOMMENDED"),
        ("ats_max", "ATS-MAX VARIANT"),
        ("creative", "CREATIVE VARIANT"),
    ]:
        content = results.get('cv_variants', {}).get(variant_key, '')
        if content:
            lines += ["", sep, variant_label, sep, "", content]

    cl = results.get('cover_letter', '')
    if cl:
        lines += ["", sep, "COVER LETTER", sep, "", cl]

    lines += ["", sep, "DETAILED AGENT REPORTS", sep]
    for name, data in results.get('agent_results', {}).items():
        lines += [
            "",
            f"─── {name.replace('_', ' ').upper()} — {data.get('score', 0)}/100 ───",
            "Findings:",
        ]
        for f in data.get('findings', []):
            lines.append(f"  • {f}")
        lines.append("Recommendations:")
        for r in data.get('recommendations', []):
            lines.append(f"  → {r}")
        if data.get('optimized_content'):
            lines += ["AI Improvement:", data['optimized_content']]

    return '\n'.join(lines)
