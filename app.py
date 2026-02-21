"""
ATS-GOD v4.0 — Main Streamlit Application
Fully wired: uploads → 9 agents → 3 CV variants + cover letter + export
Supports Groq (free) + OpenAI + Anthropic
Copyright © 2025 [Your Name/Company]
All rights reserved.

Unauthorized copying, distribution, or modification of this software,
via any medium, is strictly prohibited without express written permission.
"""
import asyncio
import logging
import os
import sys
from io import BytesIO
from datetime import datetime

import streamlit as st
from dotenv import load_dotenv

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Page Config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ATS-GOD | CV Optimizer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""<style>
    .main-header{font-size:2.8rem;font-weight:900;color:#1565C0;text-align:center;margin-bottom:0}
    .sub-header{font-size:1.05rem;color:#666;text-align:center;margin-bottom:1.5rem}
    .score-card{background:linear-gradient(135deg,#1565C0,#0D47A1);color:white;
        border-radius:14px;padding:22px;text-align:center;margin-bottom:10px}
    .score-number{font-size:3.8rem;font-weight:900;line-height:1}
    .agent-block{border-left:4px solid #1565C0;padding:8px 14px;
        margin:5px 0;background:#f7f9ff;border-radius:0 8px 8px 0}
    .green{color:#2E7D32;font-weight:bold}
    .orange{color:#E65C00;font-weight:bold}
    .red{color:#C62828;font-weight:bold}
    .pill-green{background:#E8F5E9;color:#2E7D32;padding:3px 10px;border-radius:20px;font-size:.82rem}
    .pill-orange{background:#FFF3E0;color:#E65C00;padding:3px 10px;border-radius:20px;font-size:.82rem}
    .pill-red{background:#FFEBEE;color:#C62828;padding:3px 10px;border-radius:20px;font-size:.82rem}
    .stButton>button{font-size:1.05rem;font-weight:700;padding:14px}
</style>""", unsafe_allow_html=True)

# ─── File Extraction ──────────────────────────────────────────────────────────

def read_pdf(file) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        if len(text.strip()) < 100:
            raise ValueError("Too little text extracted")
        return text
    except Exception:
        try:
            import pdfplumber
            with pdfplumber.open(file) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception as e:
            st.error(f"PDF read failed: {e}. Try pasting text instead.")
            return ""


def read_docx(file) -> str:
    try:
        from docx import Document
        doc = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts)
    except Exception as e:
        st.error(f"DOCX read failed: {e}")
        return ""

# ─── API Key Check ────────────────────────────────────────────────────────────

def detect_available_llm():
    checks = [
        ("GROQ_API_KEY", "gsk_your", "Groq (free)"),
        ("OPENAI_API_KEY", "sk-your", "OpenAI"),
        ("ANTHROPIC_API_KEY", "sk-ant-your", "Anthropic"),
    ]
    for env_var, placeholder, label in checks:
        val = os.getenv(env_var, "")
        if val and not val.startswith(placeholder):
            return True, label, "active"
    return False, None, "none"

# ─── Sidebar ──────────────────────────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.header("⚙️ Configuration")

        has_key, provider, status = detect_available_llm()
        if has_key:
            st.success(f"🟢 {provider} Connected")
        else:
            st.error("🔴 No API Key Found")
            st.info("""**Get free AI (Groq):**
1. Go to console.groq.com
2. Sign up free (no credit card)
3. Create API key
4. Add to GitHub Secrets as `GROQ_API_KEY`
5. Restart Codespace""")

        st.divider()
        st.subheader("🎯 Optimization Settings")

        target_market = st.selectbox(
            "Target Market",
            ["Both", "South Africa", "International"],
            help="Weights agent priorities accordingly"
        )
        experience_level = st.selectbox(
            "Experience Level",
            ["Entry (0-2 yrs)", "Mid (3-7 yrs)", "Senior (8-15 yrs)", "Executive (15+ yrs)"],
            index=1
        )
        industry = st.text_input("Industry", placeholder="e.g. FinTech, Mining, Healthcare")
        target_role = st.text_input("Target Role", placeholder="e.g. Senior Data Analyst")
        generate_cl = st.checkbox("✉️ Generate Cover Letter", value=True)

        st.divider()
        st.caption("**ATS-GOD v4.0** | 9 AI agents")
        st.caption("Groq (free) · OpenAI · Anthropic")

        return {
            "target_market": target_market,
            "experience_level": experience_level.split(" ")[0],
            "industry": industry,
            "target_role": target_role,
            "generate_cover_letter": generate_cl,
        }

# ─── Results ──────────────────────────────────────────────────────────────────

def render_results(results: dict):
    summary = results.get("summary", {})
    overall = summary.get("overall_score", 0)
    prob = summary.get("interview_probability", 0)

    # Hero metrics
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""<div class="score-card">
            <div style="font-size:.85rem;opacity:.8;">ATS SCORE</div>
            <div class="score-number">{overall}</div>
            <div style="font-size:.85rem;opacity:.8;">/100</div>
        </div>""", unsafe_allow_html=True)

    with c2:
        delta = f"+{prob-50}% vs avg" if prob > 50 else f"{prob-50}% vs avg"
        st.metric("Interview Probability", f"{prob}%", delta)
        st.metric("Recommended Variant", summary.get("recommended_variant", "BALANCED"))

    with c3:
        provider = results.get("llm_provider", "Rule-Based")
        model = results.get("llm_model", "N/A")
        time_s = results.get("metadata", {}).get("execution_seconds", 0)
        st.metric("AI Engine", provider)
        st.metric("Analysis Time", f"{time_s}s")
        st.caption(f"Model: {model}")

    st.info(f"📋 **Verdict:** {summary.get('verdict', '')}")

    # Agent score grid
    st.subheader("🤖 Agent Analysis")
    ICONS = {
        "algorithm_breaker": "🎯", "sa_specialist": "🇿🇦", "global_setter": "🌍",
        "recruiter_scanner": "👁️", "hiring_manager": "💼", "semantic_matcher": "📊",
        "compliance_guardian": "⚖️", "future_architect": "🚀",
    }

    cols = st.columns(4)
    for i, (name, s) in enumerate(summary.get("agent_scores", {}).items()):
        with cols[i % 4]:
            icon = ICONS.get(name, "🤖")
            label = name.replace("_", " ").title()
            color = "#2E7D32" if s >= 80 else "#E65C00" if s >= 60 else "#C62828"
            bar = "█" * (s // 10) + "░" * (10 - s // 10)
            st.markdown(f"""<div class="agent-block">
                <b>{icon} {label}</b><br>
                <span style="color:{color};font-size:1.3rem;font-weight:bold">{s}</span>/100<br>
                <span style="font-size:.75rem;color:#888">{bar}</span>
            </div>""", unsafe_allow_html=True)

    # Action Items
    st.subheader("✅ Priority Action Items")
    items = results.get("action_items", [])
    if items:
        for i, item in enumerate(items, 1):
            icon = "🔴" if i <= 4 else "🟡" if i <= 8 else "🟢"
            st.markdown(f"{icon} **{i}.** {item}")
    else:
        st.info("No action items generated — try adding an API key for AI analysis")

    # CV Variants
    st.subheader("📄 Optimized CV Variants")
    t1, t2, t3 = st.tabs(["⚡ ATS-MAX", "⚖️ BALANCED ⭐", "🎨 CREATIVE"])
    variants = results.get("cv_variants", {})

    with t1:
        st.info("Maximum ATS score. For large corporations and strict ATS portals.")
        st.text_area("ATS-MAX", variants.get("ats_max", ""), height=450, key="v_ats")
    with t2:
        st.success("Best of both worlds. Recommended for 90% of applications.")
        st.text_area("BALANCED", variants.get("balanced", ""), height=450, key="v_bal")
    with t3:
        st.info("Human-first. Best for startups, agencies, creative roles.")
        st.text_area("CREATIVE", variants.get("creative", ""), height=450, key="v_cre")

    # Cover Letter
    cl = results.get("cover_letter", "")
    if cl:
        st.subheader("✉️ Cover Letter")
        st.text_area("Cover Letter (customize before sending)", cl, height=350, key="cl")

    # Detailed agent reports (expandable)
    with st.expander("🔍 Full Agent Reports (click to expand)"):
        for name, data in results.get("agent_results", {}).items():
            icon = ICONS.get(name, "🤖")
            s = data.get("score", 0)
            color = "green" if s >= 80 else "orange" if s >= 60 else "red"
            with st.expander(f"{icon} {name.replace('_',' ').title()} — {s}/100"):
                st.write("**Findings:**")
                for f in data.get("findings", []):
                    st.write(f"• {f}")
                st.write("**Recommendations:**")
                for r in data.get("recommendations", []):
                    st.write(f"→ {r}")
                if data.get("optimized_content"):
                    st.success(f"💡 **AI Improvement:** {data['optimized_content']}")

    # Downloads
    st.subheader("💾 Download Results")
    d1, d2 = st.columns(2)

    with d1:
        from src.core.exporter import export_to_txt
        txt = export_to_txt(results)
        st.download_button(
            "📄 Download Full Report (TXT)",
            data=txt,
            file_name=f"ats_god_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
            mime="text/plain",
            use_container_width=True
        )

    with d2:
        try:
            from src.core.exporter import export_to_docx
            docx = export_to_docx(results)
            if docx:
                st.download_button(
                    "📝 Download Report (DOCX)",
                    data=docx,
                    file_name=f"ats_god_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.info("DOCX export unavailable — python-docx not installed")
        except Exception:
            st.info("DOCX export: run `pip install python-docx`")

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    st.markdown('<h1 class="main-header">🤖 ATS-GOD</h1>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sub-header">8-Agent CV Optimizer · South Africa + International · Groq / OpenAI</p>',
        unsafe_allow_html=True
    )

    context = render_sidebar()

    # Inputs
    col_cv, col_jd = st.columns(2)

    with col_cv:
        st.subheader("📤 Your CV")
        method = st.radio("Input method", ["Upload PDF/DOCX", "Paste text"], horizontal=True, key="method")
        cv_text = ""

        if method == "Upload PDF/DOCX":
            uploaded = st.file_uploader("Upload CV", type=["pdf", "docx"], key="upload")
            if uploaded:
                size = len(uploaded.getvalue()) / (1024 * 1024)
                if size > 10:
                    st.error("Max file size: 10MB")
                else:
                    buf = BytesIO(uploaded.read())
                    with st.spinner("Extracting text..."):
                        cv_text = read_pdf(buf) if "pdf" in uploaded.type else read_docx(buf)
                    if cv_text:
                        wc = len(cv_text.split())
                        st.success(f"✓ {wc} words extracted from {uploaded.name}")
                        if wc < 100:
                            st.warning("Very short CV — is it a scanned image? Try pasting text.")
                        with st.expander("Preview"):
                            st.text(cv_text[:1500] + ("…" if len(cv_text) > 1500 else ""))
        else:
            cv_text = st.text_area(
                "Paste your CV",
                height=380,
                placeholder="Copy and paste your complete CV here…",
                key="cv_paste"
            )

    with col_jd:
        st.subheader("📋 Job Description")
        jd_text = st.text_area(
            "Paste the job description",
            height=420,
            placeholder="Copy and paste the complete job description here…",
            key="jd"
        )

    st.divider()

    # Validate
    cv_ok = bool(cv_text and cv_text.strip())
    jd_ok = bool(jd_text and jd_text.strip())

    if not cv_ok:
        st.warning("👆 Add your CV — upload a file or paste text")
    if not jd_ok:
        st.warning("👆 Paste the job description")

    if cv_ok and jd_ok:
        wc_cv = len(cv_text.split())
        wc_jd = len(jd_text.split())

        m1, m2, m3 = st.columns(3)
        m1.metric("CV Word Count", wc_cv,
                  "✓ Good" if 300 <= wc_cv <= 1200 else ("Too short" if wc_cv < 300 else "May be too long"))
        m2.metric("JD Word Count", wc_jd,
                  "✓ Detailed" if wc_jd >= 120 else "More detail = better analysis")
        _, _, label = detect_available_llm()
        m3.metric("AI Mode", label or "Rule-Based",
                  "Full AI analysis" if label else "Add GROQ_API_KEY for AI")

        if st.button("🚀 OPTIMIZE — LAUNCH ALL 8 AGENTS", type="primary", use_container_width=True):

            prog = st.progress(0.0)
            status = st.empty()

            def cb(pct: float, msg: str):
                prog.progress(min(pct, 1.0))
                status.text(msg)

            try:
                sys.path.insert(0, ".")
                from src.core.orchestrator import ATSGodOrchestrator

                orchestrator = ATSGodOrchestrator()

                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

                results = loop.run_until_complete(
                    orchestrator.optimize(
                        cv_text=cv_text,
                        job_description=jd_text,
                        context=context,
                        generate_cover_letter=context.get("generate_cover_letter", True),
                        progress_callback=cb,
                    )
                )
                loop.close()

                prog.progress(1.0)
                status.text("✅ All agents complete!")
                st.balloons()

                elapsed = results.get("metadata", {}).get("execution_seconds", 0)
                st.success(f"✅ Optimization complete in {elapsed}s — {results.get('llm_provider', 'Rule-Based')} mode")
                st.divider()
                render_results(results)

            except Exception as e:
                st.error(f"❌ Optimization failed: {e}")
                import traceback
                with st.expander("Debug info"):
                    st.code(traceback.format_exc())
                st.info("Check: Are all files in src/ present? Is PYTHONPATH set? Run: `export PYTHONPATH=.`")


if __name__ == "__main__":
    main()

st.sidebar.markdown("---")
st.sidebar.caption("© 2025 ATS-GOD™ | All rights reserved")
st.sidebar.caption("Proprietary software — Unauthorized use prohibited")
